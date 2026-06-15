"""生成《Python数据分析与展示》项目实践报告"""
import glob
import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BASE_DIR = r"I:\class_project\python"
TEMPLATE_GLOB = os.path.join(BASE_DIR, "*项目实践报告模板*.doc")
OUTPUT_DOCX = os.path.join(BASE_DIR, "Python数据分析与展示_项目实践报告（2026春）.docx")
OUTPUT_DOC = os.path.join(BASE_DIR, "Python数据分析与展示_项目实践报告（2026春）.doc")


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(24)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    return p


def add_image(doc, filename, caption):
    path = os.path.join(BASE_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        set_run_font(run, size=10.5)


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《Python数据分析与展示》\n项目实践报告")
    set_run_font(run, size=16, bold=True)

    cover_lines = [
        "报告名称：杭州市租房市场数据采集、分析与可视化",
        "专业名称：计算机科学与技术",
        "班    级：请填写班级",
        "小组组长：请填写姓名（学号）",
        "成    员：请填写姓名（学号）",
        "成    员：请填写姓名（学号）",
        "指导教师：请填写教师姓名",
    ]
    for line in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(line), size=12)

    doc.add_page_break()

    add_heading(doc, "1 引言")
    add_heading(doc, "1.1 研究背景与意义", level=2)
    add_body(
        doc,
        "随着城市化进程加快，杭州作为新一线城市，租房需求持续增长。租房者在选择房源时，"
        "往往需要从大量信息中比较行政区、商圈、面积与租金之间的关系。贝壳找房等平台虽然提供了"
        "丰富的房源信息，但数据分散在网页中，难以直接用于统计分析。本项目以杭州市租房市场为研究对象，"
        "通过网络爬虫获取公开列表数据，并结合数据清洗、统计分析、聚类挖掘与地图可视化方法，"
        "探索不同区域租金水平、面积与单价之间的关系，为租房决策和区域对比提供数据支持。"
        "该实践对于理解 Python 数据分析完整流程、掌握反爬处理与可视化展示方法具有较强现实意义。"
    )

    add_heading(doc, "1.2 研究目标与主要内容", level=2)
    add_body(
        doc,
        "本项目目标包括：（1）从贝壳找房杭州租房频道抓取房源标题、行政区、商圈、面积和月租金；"
        "（2）对原始数据进行缺失值处理、重复值删除、类型转换和异常值剔除；"
        "（3）按行政区统计平均租金、平均面积和平均单价，并绘制箱线图与回归散点图；"
        "（4）基于面积、租金和单价进行 KMeans 聚类，识别不同性价比房源群体；"
        "（5）使用 pyecharts 绘制杭州市域租金单价热力地图，实现空间可视化展示。"
    )

    add_heading(doc, "1.3 开发环境", level=2)
    add_body(
        doc,
        "操作系统：Windows 10；开发工具：Jupyter Notebook / Cursor；"
        "编程语言：Python 3.12；主要依赖库：requests、lxml、pandas、matplotlib、seaborn、"
        "scikit-learn、pyecharts；数据来源：https://hz.zu.ke.com/zufang/；"
        "项目代码仓库：https://github.com/Ckj6818/hangzhou-rent-analysis。"
    )

    add_heading(doc, "2 需求分析")
    add_heading(doc, "2.1 分析需求说明", level=2)
    add_body(
        doc,
        "业务需求围绕“帮助用户快速了解杭州各区域租房价格特征”展开。具体包括："
        "获取多页租房列表数据；提取标题、行政区、商圈、面积、租金等核心字段；"
        "清洗价格与面积中的中文单位；计算每平米单价；比较不同行政区的租金分布；"
        "分析面积与租金的相关性；通过聚类识别高单价、低单价及大面积低租金等典型群体；"
        "最终在地图上用颜色深浅展示各区域平均单价差异。"
    )

    add_heading(doc, "2.2 相关 Python 包及说明", level=2)
    add_body(
        doc,
        "requests 用于发送 HTTP 请求获取网页源码；lxml 的 etree 模块用于 XPath 解析网页结构；"
        "pandas 用于数据存储、清洗和分组统计；matplotlib 与 seaborn 用于箱线图、散点图和三维散点图绘制；"
        "scikit-learn 提供 StandardScaler、KMeans 和 silhouette_score，用于标准化、聚类及模型评价；"
        "pyecharts 的 Map 组件用于生成交互式杭州市热力地图。"
    )

    add_heading(doc, "3 数据获取")
    add_heading(doc, "3.1 数据获取方式", level=2)
    add_body(
        doc,
        "本项目采用 requests + lxml 的方式抓取贝壳找房杭州租房列表页。"
        "列表页 URL 规则为：第 1 页 https://hz.zu.ke.com/zufang/ ，"
        "第 n 页 https://hz.zu.ke.com/zufang/pgn/ 。"
        "爬虫设置了 5 个常见浏览器 User-Agent 池，每次请求随机选择，并在翻页之间加入 1.5~3.5 秒随机延时。"
        "实际运行中，第 1 页成功解析 30 条房源；第 2 页起触发平台风控时，程序会重试并保留已成功抓取的数据。"
        "最终原始数据保存为 hangzhou_rent_raw.csv。"
    )

    add_heading(doc, "3.2 关键源码", level=2)
    add_code(
        doc,
        "items = tree.xpath(\"//div[@class='content__list--item']\")\n"
        "title = item.xpath(\".//p[contains(@class,'title')]//a/text()\")\n"
        "district = item.xpath(\".//p[contains(@class,'des')]/a[1]/text()\")\n"
        "price = item.xpath(\".//span[contains(@class,'price')]//em/text()\")"
    )

    add_heading(doc, "3.3 数据介绍", level=2)
    add_body(
        doc,
        "原始数据共 30 条，字段包括 title（房源标题）、district（行政区）、biz_circle（商圈）、"
        "area（面积）、price（月租金）。样本覆盖拱墅区、余杭区、钱塘区、萧山区、西湖区、上城区、临安区等。"
        "其中部分条目为独栋/公寓类房源，存在行政区缺失或价格区间表达（如 3850-4500）的情况，"
        "需要在预处理阶段进一步清洗。"
    )

    add_heading(doc, "4 数据预处理")
    add_heading(doc, "4.1 缺失值与重复值处理", level=2)
    add_heading(doc, "4.1.1 预处理目标", level=2)
    add_body(doc, "删除包含缺失值的记录和完全重复的记录，保证后续分析字段完整。")
    add_heading(doc, "4.1.2 关键源码", level=2)
    add_code(doc, "df_clean = df_raw.dropna().drop_duplicates().copy()")
    add_heading(doc, "4.1.3 预处理结果", level=2)
    add_body(doc, "原始 30 条数据中，经缺失值和重复值处理后保留有效记录，为后续类型转换做准备。")

    add_heading(doc, "4.2 类型转换与衍生指标", level=2)
    add_heading(doc, "4.2.1 预处理目标", level=2)
    add_body(doc, "将 area、price 中的“平米”“元/月”等中文及单位剥离，并计算 unit_price = price / area。")
    add_heading(doc, "4.2.2 关键源码", level=2)
    add_code(
        doc,
        "numeric_text = re.sub(r\"[^\\d.]\", \"\", text)\n"
        "df_clean['unit_price'] = (df_clean['price'] / df_clean['area']).round(2)"
    )
    add_heading(doc, "4.2.3 预处理结果", level=2)
    add_body(doc, "area 与 price 成功转换为 float 类型，并新增 unit_price 字段。")

    add_heading(doc, "4.3 IQR 异常值剔除", level=2)
    add_heading(doc, "4.3.1 预处理目标", level=2)
    add_body(doc, "采用四分位距法（IQR）分别对 area 和 price 进行异常值检测，剔除极端样本。")
    add_heading(doc, "4.3.2 关键源码", level=2)
    add_code(
        doc,
        "lower = q1 - 1.5 * iqr\nupper = q3 + 1.5 * iqr\n"
        "df = df[(df[column] >= lower) & (df[column] <= upper)]"
    )
    add_heading(doc, "4.3.3 预处理结果", level=2)
    add_body(
        doc,
        "最终清洗后数据保存为 hangzhou_rent_clean.csv，共 16 条有效样本。"
        "与原始 30 条相比，主要剔除了行政区缺失、价格区间无法直接转换及极端面积/租金记录。"
    )

    add_heading(doc, "5 统计分析及结论")
    add_heading(doc, "5.1 各行政区租金指标对比", level=2)
    add_heading(doc, "5.1.1 分析目标", level=2)
    add_body(doc, "按行政区统计平均租金、平均面积和平均单价，比较区域差异。")
    add_heading(doc, "5.1.2 关键源码", level=2)
    add_code(
        doc,
        "district_stats = df.groupby('district').agg(\n"
        "    平均租金=('price','mean'), 平均面积=('area','mean'), 平均单价=('unit_price','mean'))"
    )
    add_heading(doc, "5.1.3 分析结论", level=2)
    add_body(
        doc,
        "统计结果显示：上城区平均租金 3550 元/月、平均单价 60.36 元/㎡，属于较高水平；"
        "临安区平均租金 1800 元/月、平均单价 14.40 元/㎡，整体价格较低；"
        "萧山区样本数最多（4 条），平均单价 46.92 元/㎡；余杭区平均面积最大，达到 105㎡，"
        "但平均单价仅 30.43 元/㎡，呈现“面积大、总租金中等、单价偏低”的特征。"
    )

    add_heading(doc, "5.2 面积与租金相关性分析", level=2)
    add_heading(doc, "5.2.1 分析目标", level=2)
    add_body(doc, "分析房屋面积与月租金之间是否存在线性关系。")
    add_heading(doc, "5.2.2 关键源码", level=2)
    add_code(doc, "sns.regplot(data=df, x='area', y='price')")
    add_heading(doc, "5.2.3 分析结论", level=2)
    add_body(
        doc,
        "面积与租金的 Pearson 相关系数约为 0.45，呈中等强度正相关。"
        "说明面积越大，月租金总体越高，但租金还明显受到行政区、商圈、房源类型等因素影响，"
        "因此部分大面积房源并不一定对应最高租金。"
    )
    add_image(doc, "output_boxplot.png", "图1 杭州各行政区租金分布箱线图")
    add_image(doc, "output_scatter.png", "图2 杭州租房面积与租金相关性散点图")

    add_heading(doc, "6 数据挖掘分析")
    add_heading(doc, "6.1 KMeans 聚类分析", level=2)
    add_heading(doc, "6.1.1 数据挖掘目标", level=2)
    add_body(doc, "以 area、price、unit_price 为特征，对房源进行聚类，挖掘不同性价比群体。")
    add_heading(doc, "6.1.2 关键源码", level=2)
    add_code(
        doc,
        "X_scaled = StandardScaler().fit_transform(df[['area','price','unit_price']])\n"
        "for k in range(2, 7):\n    score = silhouette_score(X_scaled, KMeans(k).fit_predict(X_scaled))"
    )
    add_heading(doc, "6.1.3 建模结论", level=2)
    add_body(
        doc,
        "对 K=2~6 的轮廓系数进行比较，K=6 时轮廓系数最高（约 0.575）。"
        "聚类结果显示：部分簇表现为“大面积 + 低单价”，如临安、闲林等区域样本；"
        "部分簇表现为“小面积 + 高单价”，如上城采荷、萧山合租样本；"
        "也有中间价位、面积适中的主流租赁群体。该结果可为不同预算租客提供分群参考。"
    )
    add_image(doc, "output_cluster_3d.png", "图3 杭州租房 KMeans 聚类三维分布图")

    add_heading(doc, "6.2 模型评价与调优", level=2)
    add_heading(doc, "6.2.1 数据挖掘目标", level=2)
    add_body(doc, "通过轮廓系数选择较优聚类数，并讨论样本量对聚类效果的影响。")
    add_heading(doc, "6.2.2 关键源码", level=2)
    add_code(doc, "best_k = silhouette_df.loc[silhouette_df['silhouette_score'].idxmax(), 'K']")
    add_heading(doc, "6.2.3 建模结论", level=2)
    add_body(
        doc,
        "轮廓系数随 K 值变化显示模型在 K=6 时最优，但由于当前有效样本仅 16 条，"
        "聚类结果更适合作为方法演示，若用于答辩展示或实际决策，建议扩大抓取页数、"
        "补充 Cookie 后重新采集，以获得更稳定的分群结果。"
    )

    add_heading(doc, "7 数据可视化展示")
    add_heading(doc, "7.1 行政区租金分布箱线图", level=2)
    add_heading(doc, "7.1.1 可视化展示目标", level=2)
    add_body(doc, "展示不同行政区租金分布差异，识别高租金区域。")
    add_heading(doc, "7.1.2 关键源码", level=2)
    add_code(doc, "sns.boxplot(data=df, x='district', y='price')")
    add_heading(doc, "7.1.3 可视化展示效果及结论", level=2)
    add_body(doc, "箱线图显示上城、西湖等核心区域租金水平整体偏高，临安、萧山部分样本租金相对较低。")

    add_heading(doc, "7.2 杭州市租金单价热力地图", level=2)
    add_heading(doc, "7.2.1 可视化展示目标", level=2)
    add_body(doc, "在杭州市地图上展示各行政区平均租金单价，形成空间可视化效果。")
    add_heading(doc, "7.2.2 关键源码", level=2)
    add_code(
        doc,
        "Map().add('平均租金单价', map_data, maptype='杭州')\n"
        ".set_global_opts(visualmap_opts=VisualMapOpts(...), tooltip_opts=TooltipOpts(...))"
    )
    add_heading(doc, "7.2.3 可视化展示效果及结论", level=2)
    add_body(
        doc,
        "最终生成 hangzhou_rent_map.html 交互式地图。地图支持鼠标悬停查看各区平均单价，"
        "VisualMap 颜色深浅反映单价高低。结果显示：上城区、拱墅区、西湖区颜色较深，"
        "临安区、余杭部分区域颜色较浅，直观体现了杭州租房市场的空间差异。"
        "该地图适合作为答辩压轴展示内容。"
    )

    add_heading(doc, "8 分析结论及应用方向")
    add_heading(doc, "8.1 分析总结", level=2)
    add_body(
        doc,
        "（1）本项目完成了从数据采集、清洗、统计分析、聚类挖掘到地图可视化的完整流程；"
        "（2）杭州各区域租金差异明显，核心城区单价普遍高于临安、余杭部分区域；"
        "（3）面积与租金呈中等正相关，但房源位置和类型对租金影响显著；"
        "（4）KMeans 聚类可初步识别高单价、低单价及大面积低租金等不同群体；"
        "（5）pyecharts 热力地图能直观展示空间分布，是较有展示性的分析成果。"
    )

    add_heading(doc, "8.2 应用方向", level=2)
    add_body(
        doc,
        "本项目的分析结论可应用于：为初到杭州的租客提供区域租金参考；"
        "为房产平台运营人员提供区域定价对比；为城市租赁市场研究提供数据样本；"
        "后续还可结合地铁线路、学区、房龄等变量进行多元回归或推荐系统建模。"
    )

    add_heading(doc, "结束语")
    add_body(
        doc,
        "设计过程中遇到的主要问题包括：贝壳网站反爬导致第 2 页以后请求被重定向到登录页、"
        "部分房源价格以区间形式展示、个别记录缺少行政区字段。"
        "解决方法包括：加入 Cookie、固定 User-Agent、增加重试与延时、对失败页保留已成功数据、"
        "在清洗阶段删除无法转换的记录并使用 IQR 剔除异常值。"
        "通过本次实践，完整掌握了 Python 数据分析项目从需求分析到可视化交付的流程，"
        "也认识到真实数据采集比实验数据更复杂，数据质量直接决定分析结论的可信度。"
    )

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def convert_to_doc(docx_path, doc_path):
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    doc.SaveAs2(doc_path, FileFormat=0)
    doc.Close(False)
    word.Quit()


if __name__ == "__main__":
    docx_path = build_document()
    convert_to_doc(docx_path, OUTPUT_DOC)
    print(f"报告已生成: {OUTPUT_DOC}")
